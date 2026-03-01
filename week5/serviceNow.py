import time
import os
import re
import json
from urllib.parse import urljoin, urlparse
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
from fpdf import FPDF
from docx import Document
import requests
from bs4 import BeautifulSoup

class ServiceNowDocScraper:
    def __init__(self, output_format='pdf', delay=5):
        """
        Initialize the scraper with enhanced anti-detection measures
        :param output_format: 'pdf' or 'docx'
        :param delay: Delay between requests in seconds
        """
        self.output_format = output_format
        self.delay = delay
        self.driver = None
        self.base_url = "https://www.servicenow.com"
        self.session = requests.Session()
        
    def setup_driver(self):
        """Setup Chrome driver with enhanced stealth options"""
        chrome_options = Options()
        
        # Enhanced stealth options
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)
        chrome_options.add_argument("--window-size=1920,1080")
        
        # Realistic user agent
        chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
        
        # Additional stealth options
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--disable-plugins-discovery")
        chrome_options.add_argument("--disable-web-security")
        chrome_options.add_argument("--allow-running-insecure-content")
        chrome_options.add_argument("--no-first-run")
        chrome_options.add_argument("--disable-default-apps")
        
        # Set up service
        service = Service(ChromeDriverManager().install())
        self.driver = webdriver.Chrome(service=service, options=chrome_options)
        
        # Execute script to hide webdriver property
        self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        # Set realistic viewport and headers
        self.driver.execute_cdp_cmd('Network.setUserAgentOverride', {
            "userAgent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        })
        
    def handle_authentication(self):
        """Handle potential authentication or access requirements"""
        try:
            # Check if we need to accept cookies or terms
            cookie_selectors = [
                'button[id*="cookie"]',
                'button[class*="cookie"]',
                'button[id*="accept"]',
                'button[class*="accept"]',
                '.cookie-accept',
                '#cookie-accept',
                '.accept-cookies'
            ]
            
            for selector in cookie_selectors:
                try:
                    button = WebDriverWait(self.driver, 3).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, selector))
                    )
                    button.click()
                    print("Accepted cookies/terms")
                    time.sleep(2)
                    break
                except:
                    continue
                    
        except Exception as e:
            print(f"No cookie banner found or couldn't handle: {e}")
    
    def get_documentation_links(self, main_url):
        """Extract all documentation links with enhanced error handling"""
        print(f"Fetching main page: {main_url}")
        
        try:
            # Navigate to the main page
            self.driver.get(main_url)
            print("Page loaded, waiting for content...")
            
            # Handle any authentication/cookie banners
            self.handle_authentication()
            
            # Wait longer for content to load
            time.sleep(10)
            
            # Check if we got an error page
            page_source = self.driver.page_source.lower()
            if any(error in page_source for error in ['error', 'not found', 'access denied', 'protocol_error']):
                print("Error page detected. Trying alternative approach...")
                
                # Try to find any accessible documentation links in the current page
                self.try_alternative_link_extraction()
            
            # Wait for specific content indicators
            content_indicators = [
                "//a[contains(@href, '/docs/')]",
                "//a[contains(text(), 'documentation')]",
                "//a[contains(text(), 'guide')]",
                "//a[contains(text(), 'reference')]",
                "//div[contains(@class, 'content')]",
                "//main",
                "//article"
            ]
            
            found_content = False
            for indicator in content_indicators:
                try:
                    elements = self.driver.find_elements(By.XPATH, indicator)
                    if elements:
                        found_content = True
                        print(f"Found content with indicator: {indicator}")
                        break
                except:
                    continue
            
            if not found_content:
                print("No content indicators found. Page might be blocked or require authentication.")
                return self.get_fallback_links()
            
            # Extract links
            links = []
            link_elements = self.driver.find_elements(By.TAG_NAME, "a")
            
            print(f"Found {len(link_elements)} total links on page")
            
            for element in link_elements:
                try:
                    href = element.get_attribute("href")
                    text = element.text.strip()
                    
                    if href and text and len(text) > 3:
                        # Filter for documentation-related links
                        if any(keyword in href.lower() for keyword in ["/docs/", "documentation", "reference", "guide", "bundle"]):
                            # Make sure it's a full URL
                            if href.startswith('/'):
                                href = urljoin(self.base_url, href)
                            
                            links.append({
                                'url': href,
                                'title': text,
                                'clean_title': self.clean_filename(text)
                            })
                except Exception as e:
                    continue
            
            # Remove duplicates
            unique_links = []
            seen_urls = set()
            for link in links:
                if link['url'] not in seen_urls:
                    unique_links.append(link)
                    seen_urls.add(link['url'])
            
            print(f"Found {len(unique_links)} unique documentation links")
            
            if len(unique_links) == 0:
                print("No documentation links found. Trying fallback method...")
                return self.get_fallback_links()
            
            return unique_links
            
        except Exception as e:
            print(f"Error fetching main page: {e}")
            return self.get_fallback_links()
    
    def try_alternative_link_extraction(self):
        """Try alternative methods to extract links"""
        try:
            # Try JavaScript execution to get links
            js_script = """
            var links = [];
            var anchors = document.getElementsByTagName('a');
            for(var i = 0; i < anchors.length; i++) {
                var href = anchors[i].href;
                var text = anchors[i].innerText;
                if(href && text && (href.includes('/docs/') || href.includes('documentation'))) {
                    links.push({url: href, title: text});
                }
            }
            return links;
            """
            
            js_links = self.driver.execute_script(js_script)
            print(f"JavaScript extraction found {len(js_links)} links")
            return js_links
            
        except Exception as e:
            print(f"JavaScript extraction failed: {e}")
            return []
    
    def get_fallback_links(self):
        """Provide fallback documentation links if main page is inaccessible"""
        print("Using fallback documentation links...")
        
        # Common ServiceNow documentation URLs
        fallback_links = [
            {
                'url': 'https://docs.servicenow.com/bundle/yokohama-platform-administration/page/administer/overview.html',
                'title': 'Platform Administration Guide',
                'clean_title': 'Platform_Administration_Guide'
            },
            {
                'url': 'https://docs.servicenow.com/bundle/yokohama-application-development/page/build/overview.html',
                'title': 'Application Development Guide',
                'clean_title': 'Application_Development_Guide'
            },
            {
                'url': 'https://docs.servicenow.com/bundle/yokohama-servicenow-platform/page/use/overview.html',
                'title': 'ServiceNow Platform User Guide',
                'clean_title': 'ServiceNow_Platform_User_Guide'
            },
            {
                'url': 'https://docs.servicenow.com/bundle/yokohama-api-reference/page/integrate/overview.html',
                'title': 'API Reference Guide',
                'clean_title': 'API_Reference_Guide'
            },
            {
                'url': 'https://docs.servicenow.com/bundle/yokohama-performance-analytics/page/use/performance-analytics/concept/c_PALandingPage.html',
                'title': 'Performance Analytics Guide',
                'clean_title': 'Performance_Analytics_Guide'
            }
        ]
        
        return fallback_links
    
    def clean_filename(self, filename):
        """Clean filename for saving"""
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        filename = re.sub(r'\s+', '_', filename)
        if len(filename) > 100:
            filename = filename[:100]
        return filename
    
    def scrape_page_content(self, url):
        """Scrape content from a single page with enhanced error handling"""
        try:
            print(f"Scraping: {url}")
            self.driver.get(url)
            
            # Wait longer for content to load
            time.sleep(self.delay * 2)
            
            # Check for error pages
            page_source = self.driver.page_source.lower()
            if any(error in page_source for error in ['this site can\'t be reached', 'err_', 'error', 'access denied']):
                print(f"Error page detected for {url}")
                return {
                    'title': f"Error accessing page",
                    'content': f"Could not access {url}. Page may require authentication or be restricted.",
                    'url': url
                }
            
            # Handle authentication if needed
            self.handle_authentication()
            
            # Wait for content to load
            try:
                WebDriverWait(self.driver, 15).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
            except TimeoutException:
                print(f"Timeout waiting for {url} to load")
            
            # Get page title
            title = self.driver.title or "Untitled Document"
            
            # Enhanced content extraction
            content_selectors = [
                '.doc-content',
                '.documentation-content',
                '.main-content',
                'main',
                '.content',
                'article',
                '.article-content',
                '#content',
                '.page-content',
                '.body-content',
                '.wiki-content'
            ]
            
            content = ""
            for selector in content_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements and elements[0].text.strip():
                        content = elements[0].text
                        print(f"Content found using selector: {selector}")
                        break
                except:
                    continue
            
            # If no specific content area found, try body text
            if not content or len(content) < 100:
                try:
                    body = self.driver.find_element(By.TAG_NAME, "body")
                    body_text = body.text
                    if len(body_text) > 200:  # Only use if substantial content
                        content = body_text
                        print("Using body text as content")
                except:
                    content = "Could not extract meaningful content"
            
            # Clean up content
            content = re.sub(r'\n\s*\n', '\n\n', content)
            content = content.strip()
            
            # Validate content quality
            if len(content) < 50:
                content = f"Limited content extracted from {url}. Page may require authentication or have dynamic content."
            
            return {
                'title': title,
                'content': content,
                'url': url
            }
            
        except Exception as e:
            print(f"Error scraping {url}: {e}")
            return {
                'title': f"Error - {urlparse(url).path.split('/')[-1]}",
                'content': f"Failed to scrape content from {url}. Error: {e}",
                'url': url
            }
    
    def save_as_pdf(self, data, filename):
        """Save content as PDF with better encoding handling"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Add title
            pdf.set_font("Arial", 'B', size=16)
            title_text = data['title'][:80].encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 10, title_text, ln=True)
            pdf.ln(5)
            
            # Add URL
            pdf.set_font("Arial", size=10)
            pdf.cell(0, 10, f"Source: {data['url']}", ln=True)
            pdf.ln(5)
            
            # Add content
            pdf.set_font("Arial", size=11)
            
            lines = data['content'].split('\n')
            for line in lines[:200]:  # Limit to prevent huge files
                try:
                    line = line.encode('latin-1', 'replace').decode('latin-1')
                    # Split long lines
                    while len(line) > 120:
                        pdf.cell(0, 6, line[:120], ln=True)
                        line = line[120:]
                    if line.strip():
                        pdf.cell(0, 6, line, ln=True)
                except Exception as line_error:
                    pdf.cell(0, 6, "[Content encoding error]", ln=True)
            
            output_path = f"output/{filename}.pdf"
            os.makedirs("output", exist_ok=True)
            pdf.output(output_path)
            print(f"Saved: {output_path}")
            
        except Exception as e:
            print(f"Error saving PDF {filename}: {e}")
    
    def save_as_docx(self, data, filename):
        """Save content as DOCX with better formatting"""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading(data['title'], level=1)
            
            # Add URL
            url_para = doc.add_paragraph(f"Source: {data['url']}")
            url_para.italic = True
            
            # Add separator
            doc.add_paragraph()
            
            # Add content with better formatting
            content_lines = data['content'].split('\n')
            current_para = ""
            
            for line in content_lines[:500]:  # Limit content
                line = line.strip()
                if not line:
                    if current_para:
                        doc.add_paragraph(current_para)
                        current_para = ""
                    continue
                
                if len(current_para) + len(line) > 500:
                    doc.add_paragraph(current_para)
                    current_para = line
                else:
                    current_para += " " + line if current_para else line
            
            if current_para:
                doc.add_paragraph(current_para)
            
            output_path = f"output/{filename}.docx"
            os.makedirs("output", exist_ok=True)
            doc.save(output_path)
            print(f"Saved: {output_path}")
            
        except Exception as e:
            print(f"Error saving DOCX {filename}: {e}")
    
    def scrape_all_documentation(self, main_url):
        """Main method with enhanced error handling and recovery"""
        try:
            print("Setting up browser...")
            self.setup_driver()
            
            print("Getting documentation links...")
            doc_links = self.get_documentation_links(main_url)
            
            if not doc_links:
                print("No documentation links found! Check if the site requires authentication.")
                return
            
            print(f"Starting to scrape {len(doc_links)} documents...")
            os.makedirs("output", exist_ok=True)
            
            successful_scrapes = 0
            
            for i, link in enumerate(doc_links, 1):
                print(f"\nProgress: {i}/{len(doc_links)} - {link['title'][:50]}...")
                
                try:
                    content_data = self.scrape_page_content(link['url'])
                    filename = link['clean_title'] or f"document_{i}"
                    
                    # Save based on format
                    if self.output_format.lower() == 'pdf':
                        self.save_as_pdf(content_data, filename)
                    else:
                        self.save_as_docx(content_data, filename)
                    
                    successful_scrapes += 1
                    
                except Exception as e:
                    print(f"Failed to process {link['url']}: {e}")
                
                # Add delay between requests
                time.sleep(self.delay)
            
            print(f"\nCompleted! Successfully scraped {successful_scrapes}/{len(doc_links)} documents.")
            print(f"Files saved in 'output' directory as {self.output_format.upper()} format.")
            
            if successful_scrapes == 0:
                print("\nTroubleshooting suggestions:")
                print("1. The site may require authentication - try logging in manually first")
                print("2. ServiceNow may be blocking automated access")
                print("3. Try running with '--headless' disabled to see what's happening")
                print("4. Consider using ServiceNow's official API instead")
            
        except Exception as e:
            print(f"Error in main scraping process: {e}")
        
        finally:
            if self.driver:
                self.driver.quit()

def main():
    # URL to scrape
    main_url = "https://www.servicenow.com/docs/bundle/yokohama-product-directory/page/product-directory/reference/product-directory.html"
    
    # Choose output format
    output_format = 'pdf'  # or 'docx'
    
    # Initialize scraper with longer delays
    scraper = ServiceNowDocScraper(output_format=output_format, delay=5)
    
    print("Enhanced ServiceNow Documentation Scraper")
    print("=" * 50)
    print(f"Target URL: {main_url}")
    print(f"Output format: {output_format.upper()}")
    print(f"Delay between requests: {scraper.delay} seconds")
    print("=" * 50)
    print("\nNote: ServiceNow documentation may require authentication.")
    print("If scraping fails, consider using ServiceNow's official API or")
    print("accessing documentation through an authenticated session.")
    print("=" * 50)
    
    # Start scraping
    scraper.scrape_all_documentation(main_url)

if __name__ == "__main__":
    main()